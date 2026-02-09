"""
TrustLens AI - DOCX Document Parser
Word 文档解析模块，保留页码与段落信息
"""
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import json

# 检查并安装 python-docx
try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


from .pdf_parser import BoundingBox, TextBlock, PDFPage


class DOCXParser:
    """基于 python-docx 的 Word 文档解析器"""

    def __init__(self):
        self.engine = "python-docx"

        try:
            from docx import Document
            self.Document = Document
        except ImportError:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")

    def parse(self, file_path: str, extract_bbox: bool = True) -> List[PDFPage]:
        """
        解析 Word 文档

        Args:
            file_path: Word 文件路径
            extract_bbox: 是否提取文本块坐标（DOCX 不支持精确坐标）

        Returns:
            解析后的页面列表
        """
        pages = []

        try:
            doc = self.Document(file_path)

            # 获取文档基本信息
            # 注意：DOCX 不直接提供页码信息，需要通过段落顺序估算
            # 使用 A4 页面尺寸作为默认值（210mm x 297mm，约 595 x 842 points）
            page_width = 595.0
            page_height = 842.0

            # 收集所有段落
            all_paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue

                all_paragraphs.append({
                    'text': text,
                    'index': i,
                    'style': para.style.name if para.style else 'Normal',
                })

            # 处理表格
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if text:
                            all_paragraphs.append({
                                'text': text,
                                'index': len(all_paragraphs),
                                'style': f'Table_{table_idx}',
                                'is_table': True,
                                'table_position': (table_idx, row_idx, cell_idx),
                            })

            # 由于 DOCX 不直接提供页码，我们需要估算分页
            # 假设每页约 40 个段落（这是一个粗略估计）
            paragraphs_per_page = 40

            # 按页切分段落
            page_num = 0
            current_page_paragraphs = []

            for para_info in all_paragraphs:
                current_page_paragraphs.append(para_info)

                # 当达到每页段落数或遇到明显分页标记时创建新页
                if len(current_page_paragraphs) >= paragraphs_per_page:
                    # 创建页面
                    page = self._create_page_from_paragraphs(
                        page_num + 1,
                        page_width,
                        page_height,
                        current_page_paragraphs
                    )
                    pages.append(page)
                    current_page_paragraphs = []
                    page_num += 1

            # 添加最后一页
            if current_page_paragraphs:
                page = self._create_page_from_paragraphs(
                    page_num + 1,
                    page_width,
                    page_height,
                    current_page_paragraphs
                )
                pages.append(page)

        except Exception as e:
            raise RuntimeError(f"Word 文档解析失败: {str(e)}")

        return pages

    def _create_page_from_paragraphs(
        self,
        page_number: int,
        page_width: float,
        page_height: float,
        paragraphs: List[Dict[str, Any]]
    ) -> PDFPage:
        """从段落列表创建页面对象"""
        text_blocks = []
        full_text_lines = []

        # 垂直位置，用于模拟段落位置
        y_position = 50.0  # 起始位置
        line_height = 20.0  # 行高

        for para_info in paragraphs:
            text = para_info['text']

            # 创建边界框（模拟位置）
            bbox = BoundingBox(
                x1=50.0,  # 左边距
                y1=y_position,
                x2=page_width - 50.0,  # 右边距
                y2=y_position + line_height,
                page_width=page_width,
                page_height=page_height,
            )

            # 创建文本块
            text_block = TextBlock(
                text=text,
                page=page_number,
                bbox=bbox,
                font_name=para_info.get('style', 'Normal'),
            )

            text_blocks.append(text_block)
            full_text_lines.append(text)

            # 更新下一个段落的垂直位置
            y_position += line_height * (1 + text.count('\n'))

        # 创建页面
        page = PDFPage(
            page_number=page_number,
            width=page_width,
            height=page_height,
            text_blocks=text_blocks,
            raw_text='\n'.join(full_text_lines),
        )

        return page

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        提取 Word 文档元数据

        Args:
            file_path: Word 文件路径

        Returns:
            元数据字典
        """
        metadata = {
            "file_name": Path(file_path).name,
            "file_size": os.path.getsize(file_path),
            "pages": 0,
            "title": "",
            "author": "",
            "subject": "",
            "creator": "",
            "creation_date": "",
        }

        try:
            doc = self.Document(file_path)

            # 提取核心属性
            core_props = doc.core_properties

            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.created:
                metadata["creation_date"] = str(core_props.created)

            # 估算页数（基于段落数）
            paragraph_count = len(doc.paragraphs)
            estimated_pages = max(1, (paragraph_count + 39) // 40)  # 假设每页40段
            metadata["pages"] = estimated_pages

        except Exception as e:
            print(f"Warning: Failed to extract metadata: {e}")

        return metadata


def get_docx_parser() -> DOCXParser:
    """
    获取 DOCX 解析器实例

    Returns:
        DOCX 解析器实例
    """
    return DOCXParser()


def parse_docx(file_path: str) -> List[PDFPage]:
    """
    便捷函数：解析 Word 文档

    Args:
        file_path: Word 文件路径

    Returns:
        解析后的页面列表
    """
    parser = get_docx_parser()
    return parser.parse(file_path)
